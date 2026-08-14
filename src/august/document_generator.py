from __future__ import annotations

import os
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any

from august.answer_fallback import try_local_answer
from august.answer_memory import AnswerMemory
from august.utils.logger import get_logger, log_event

logger = get_logger("DocumentGenerator")


RELIABILITY_FAILURE_MESSAGE = "I couldn't generate a reliable document. Do you want me to search instead?"


@dataclass
class DocumentGenerationResult:
    success: bool
    message: str
    filename: str = ""
    opened: bool = False
    source: str = ""


def generate_document(
    topic: str,
    *,
    answer_memory: AnswerMemory | None = None,
    memory_store: Any | None = None,
    ai_parser: Any | None = None,
    context: dict[str, Any] | None = None,
    memory: dict[str, Any] | None = None,
    open_file: bool = True,
    output_dir: str | Path | None = None,
) -> DocumentGenerationResult:
    clean_topic = _clean_topic(topic)
    if not clean_topic:
        return DocumentGenerationResult(False, "What topic should I make the document on?")

    researched = research_topic(
        clean_topic,
        answer_memory=answer_memory,
        memory_store=memory_store,
        ai_parser=ai_parser,
        context=context,
        memory=memory,
    )
    if not researched:
        return DocumentGenerationResult(False, RELIABILITY_FAILURE_MESSAGE)

    structured = structured_formatter(clean_topic, researched["text"])
    if not _is_reliable_structure(structured):
        log_event(logger, "document_quality_rejected", source=researched["source"], success=False, topic=clean_topic)
        return DocumentGenerationResult(False, RELIABILITY_FAILURE_MESSAGE, source=researched["source"])

    try:
        filename = create_word_document(structured, output_dir=output_dir)
    except ImportError:
        logger.exception("python-docx is not available")
        return DocumentGenerationResult(False, "I need python-docx installed before I can create Word documents.", source=researched["source"])
    except Exception as exc:
        logger.exception("Failed to create document for '%s': %s", clean_topic, exc)
        return DocumentGenerationResult(False, "I couldn't save the document.", source=researched["source"])

    opened = False
    if open_file:
        try:
            os.startfile(filename)
            opened = True
        except Exception as exc:
            logger.warning("Document was created but could not be opened: %s", exc)

    open_suffix = " Opening it now." if opened else ""
    return DocumentGenerationResult(
        True,
        f"I've created a document on {clean_topic}.{open_suffix}",
        filename=filename,
        opened=opened,
        source=researched["source"],
    )


def research_topic(
    topic: str,
    *,
    answer_memory: AnswerMemory | None = None,
    memory_store: Any | None = None,
    ai_parser: Any | None = None,
    context: dict[str, Any] | None = None,
    memory: dict[str, Any] | None = None,
) -> dict[str, str] | None:
    if memory is not None:
        memory_snapshot = memory
    elif memory_store is not None and hasattr(memory_store, "snapshot"):
        memory_snapshot = memory_store.snapshot()
    else:
        memory_snapshot = {}
    query = f"what is {topic}"

    if answer_memory is None and memory_store is not None:
        answer_memory = AnswerMemory(memory_store=memory_store)

    if answer_memory is not None:
        cached = answer_memory.retrieve(query) or answer_memory.retrieve(topic)
        if _is_reliable_text(cached):
            log_event(logger, "document_research_memory_hit", source="answer_memory", success=True, topic=topic)
            return {"text": str(cached).strip(), "source": "answer_memory"}
        log_event(logger, "document_research_memory_miss", source="answer_memory", success=False, topic=topic)

    local_notes = _local_knowledge_notes(topic)
    if local_notes:
        return {"text": _sections_to_research_text(local_notes), "source": "local_knowledge"}

    local_result = try_local_answer(query)
    if local_result.get("success"):
        text = str(local_result.get("text", "")).strip()
        source = str(local_result.get("source", "local") or "local")
        if source != "local_generated" and _is_reliable_text(text):
            return {"text": text, "source": "local_knowledge"}

    if ai_parser is not None and hasattr(ai_parser, "try_ai_answer"):
        prompt = (
            f"Create accurate educational notes on {topic}. Include definition, key concepts, "
            "types or components, examples, and advantages or use cases. Use concrete facts and avoid filler."
        )
        try:
            ai_result = ai_parser.try_ai_answer(prompt, context=context or {}, memory=memory_snapshot)
        except Exception as exc:
            logger.warning("AI document research failed for '%s': %s", topic, exc)
            ai_result = {"success": False, "text": "", "error": str(exc)}
        if ai_result.get("success") and _is_reliable_text(str(ai_result.get("text", ""))):
            return {"text": str(ai_result.get("text", "")).strip(), "source": "ai"}

    return None


def structured_formatter(topic: str, research_text: str) -> dict[str, Any]:
    clean_topic = _title_case(_clean_topic(topic))
    sections = _extract_labeled_sections(research_text)
    if not sections:
        sections = _build_sections_from_plain_text(clean_topic, research_text)

    return {
        "title": clean_topic,
        "sections": [
            {"heading": heading, "content": _clean_content(content)}
            for heading, content in sections
            if _clean_content(content)
        ],
    }


def create_word_document(structured: dict[str, Any], *, output_dir: str | Path | None = None) -> str:
    from docx import Document

    title = str(structured.get("title", "")).strip() or "Notes"
    sections = structured.get("sections", [])
    directory = Path(output_dir) if output_dir is not None else Path.cwd()
    directory.mkdir(parents=True, exist_ok=True)
    filename = directory / f"{_safe_filename(title)}.docx"

    doc = Document()
    doc.add_heading(title, 0)
    for section in sections:
        doc.add_heading(str(section["heading"]), level=1)
        doc.add_paragraph(str(section["content"]))
    doc.save(str(filename))
    return str(filename)


def _local_knowledge_notes(topic: str) -> list[dict[str, str]]:
    normalized = _normalize(topic)
    if "polymorphism" in normalized:
        return [
            {
                "heading": "Definition",
                "content": "Polymorphism is an object-oriented programming concept where the same interface, method name, or operation can produce different behavior depending on the object or data type using it.",
            },
            {
                "heading": "Key Concepts",
                "content": "It allows code to work through a common contract while each class supplies its own implementation. The main idea is substitutability: a caller can use a parent type or interface without knowing the exact child class.",
            },
            {
                "heading": "Types or Components",
                "content": "Compile-time polymorphism usually appears as method overloading or operator overloading. Run-time polymorphism usually appears as method overriding, where a subclass replaces inherited behavior.",
            },
            {
                "heading": "Examples",
                "content": "A Shape interface can define draw(), while Circle, Rectangle, and Triangle each implement draw() differently. The same call, shape.draw(), runs the version that belongs to the actual object.",
            },
            {
                "heading": "Advantages / Use Cases",
                "content": "Polymorphism supports reusable code, flexible designs, plugin-style extension, cleaner conditional logic, and easier testing because new classes can be added without rewriting every caller.",
            },
        ]
    if "operating system" in normalized or "operating systems" in normalized:
        return [
            {
                "heading": "Definition",
                "content": "An operating system is system software that manages computer hardware and provides services for applications and users.",
            },
            {
                "heading": "Key Concepts",
                "content": "It coordinates processes, memory, files, devices, users, permissions, and input/output so programs can run safely and share resources.",
            },
            {
                "heading": "Types or Components",
                "content": "Major components include the kernel, process scheduler, memory manager, file system, device drivers, system calls, and user interface. Common types include batch, time-sharing, distributed, real-time, mobile, and embedded operating systems.",
            },
            {
                "heading": "Examples",
                "content": "Examples include Windows, Linux, macOS, Android, iOS, and real-time operating systems used in industrial controllers or embedded devices.",
            },
            {
                "heading": "Advantages / Use Cases",
                "content": "Operating systems make computers easier to use, isolate applications, manage limited resources, support multitasking, protect files, and provide a stable platform for software development.",
            },
        ]
    return []


def _extract_labeled_sections(text: str) -> list[tuple[str, str]]:
    wanted = {
        "definition": "Definition",
        "key concepts": "Key Concepts",
        "types or components": "Types or Components",
        "types": "Types or Components",
        "components": "Types or Components",
        "examples": "Examples",
        "advantages / use cases": "Advantages / Use Cases",
        "advantages": "Advantages / Use Cases",
        "use cases": "Advantages / Use Cases",
    }
    label_names = "definition|key concepts|types or components|types|components|examples|advantages / use cases|advantages|use cases"
    inline_pattern = re.compile(rf"(?im)^\s*(?:[-*#]+\s*)?({label_names})\s*:\s*(.+)$")
    inline_matches = list(inline_pattern.finditer(text or ""))
    if inline_matches:
        by_heading: dict[str, list[str]] = {}
        for match in inline_matches:
            heading = wanted[match.group(1).strip().lower()]
            by_heading.setdefault(heading, []).append(match.group(2).strip())
        return _complete_missing_sections([(heading, " ".join(parts)) for heading, parts in by_heading.items()])

    pattern = re.compile(rf"(?im)^\s*(?:[-*#]+\s*)?({label_names})\s*:?\s*$")
    matches = list(pattern.finditer(text or ""))
    sections: list[tuple[str, str]] = []
    seen: set[str] = set()
    for index, match in enumerate(matches):
        heading = wanted[match.group(1).strip().lower()]
        if heading in seen:
            continue
        start = match.end()
        end = matches[index + 1].start() if index + 1 < len(matches) else len(text)
        content = _clean_content(text[start:end])
        if content:
            sections.append((heading, content))
            seen.add(heading)
    return _complete_missing_sections(sections)


def _build_sections_from_plain_text(topic: str, text: str) -> list[tuple[str, str]]:
    sentences = _sentences(text)
    if not sentences:
        return []

    definition = sentences[0]
    remaining = sentences[1:]
    chunks = _chunk_sentences(remaining, 4)
    fallback = {
        "Definition": definition,
        "Key Concepts": " ".join(chunks[0]) if len(chunks) > 0 else "",
        "Types or Components": " ".join(chunks[1]) if len(chunks) > 1 else "",
        "Examples": _example_sentence(topic),
        "Advantages / Use Cases": " ".join(chunks[2]) if len(chunks) > 2 else _use_case_sentence(topic),
    }
    return [(heading, content) for heading, content in fallback.items() if content]


def _complete_missing_sections(sections: list[tuple[str, str]]) -> list[tuple[str, str]]:
    order = ["Definition", "Key Concepts", "Types or Components", "Examples", "Advantages / Use Cases"]
    by_heading = {heading: content for heading, content in sections}
    return [(heading, by_heading[heading]) for heading in order if heading in by_heading]


def _sections_to_research_text(sections: list[dict[str, str]]) -> str:
    return "\n\n".join(f"{section['heading']}\n{section['content']}" for section in sections)


def _is_reliable_structure(structured: dict[str, Any]) -> bool:
    sections = structured.get("sections")
    if not isinstance(sections, list) or len(sections) < 5:
        return False
    headings = {str(section.get("heading", "")).strip() for section in sections if isinstance(section, dict)}
    required = {"Definition", "Key Concepts", "Types or Components", "Examples", "Advantages / Use Cases"}
    if not required.issubset(headings):
        return False
    for section in sections:
        content = str(section.get("content", "")).strip()
        if len(content.split()) < 10 or _is_vague(content):
            return False
    return True


def _is_reliable_text(text: str | None) -> bool:
    clean = _clean_content(text or "")
    return len(clean.split()) >= 18 and not _is_vague(clean)


def _is_vague(text: str) -> bool:
    lowered = text.lower()
    vague_phrases = (
        "important concept",
        "core idea",
        "common use cases and trade-offs",
        "specific idea, process, or system",
        "no information",
        "i don't know",
        "unable to",
    )
    return any(phrase in lowered for phrase in vague_phrases)


def _sentences(text: str) -> list[str]:
    clean = _clean_content(text)
    return [part.strip() for part in re.split(r"(?<=[.!?])\s+", clean) if len(part.split()) >= 5]


def _chunk_sentences(sentences: list[str], size: int) -> list[list[str]]:
    return [sentences[index : index + size] for index in range(0, len(sentences), size)]


def _example_sentence(topic: str) -> str:
    return f"An example of {topic} should show the concept in a concrete program, system, or real-world situation rather than only naming it."


def _use_case_sentence(topic: str) -> str:
    return f"{topic} is useful when learners need to understand how the concept is applied, compared, and explained in practical work."


def _clean_topic(topic: str) -> str:
    cleaned = re.sub(r"\s+", " ", (topic or "").strip(" .?"))
    cleaned = re.sub(r"^(?:on|about)\s+", "", cleaned, flags=re.IGNORECASE).strip()
    return cleaned


def _clean_content(text: str) -> str:
    cleaned = re.sub(r"(?m)^\s*[-*]\s*", "", text or "")
    cleaned = re.sub(r"\s+", " ", cleaned).strip()
    return cleaned


def _normalize(text: str) -> str:
    return re.sub(r"\s+", " ", (text or "").strip().lower())


def _title_case(text: str) -> str:
    return " ".join(word.capitalize() if word.islower() else word for word in text.split())


def _safe_filename(title: str) -> str:
    cleaned = re.sub(r"[^\w\s-]", "", title).strip()
    cleaned = re.sub(r"\s+", "_", cleaned)
    return cleaned or "notes"
